"""Carga de documentos jurídicos a texto plano: .txt, .md, .docx y .pdf."""

from __future__ import annotations

from pathlib import Path


class DocumentoError(Exception):
    """Error al cargar un documento."""


class FormatoNoSoportadoError(DocumentoError):
    """Extensión no soportada por el cargador."""


class PdfEscaneadoError(DocumentoError):
    """El PDF no tiene capa de texto (parece escaneado; OCR no disponible)."""


EXTENSIONES = {".txt", ".md", ".docx", ".pdf"}


def cargar_documento(ruta: str | Path) -> str:
    """Devuelve el texto plano del documento.

    Raises:
        FormatoNoSoportadoError: si la extensión no es .txt/.md/.docx/.pdf.
        PdfEscaneadoError: si el PDF no contiene capa de texto.
        DocumentoError: si el texto resultante está vacío.
    """
    ruta = Path(ruta)
    if not ruta.is_file():
        raise DocumentoError(f"No existe el fichero: {ruta}")

    sufijo = ruta.suffix.lower()
    if sufijo in (".txt", ".md"):
        texto = ruta.read_text(encoding="utf-8")
    elif sufijo == ".docx":
        texto = _texto_docx(ruta)
    elif sufijo == ".pdf":
        texto = _texto_pdf(ruta)
    else:
        raise FormatoNoSoportadoError(
            f"Formato no soportado: {sufijo or '(sin extensión)'}. "
            f"Use: {', '.join(sorted(EXTENSIONES))}"
        )

    if not texto.strip():
        raise DocumentoError(f"El documento no contiene texto extraíble: {ruta}")
    return texto


def _texto_docx(ruta: Path) -> str:
    from docx import Document

    doc = Document(ruta)
    partes = [p.text for p in doc.paragraphs]
    for tabla in doc.tables:
        for fila in tabla.rows:
            partes.extend(celda.text for celda in fila.cells)
    return "\n".join(partes)


def _texto_pdf(ruta: Path) -> str:
    import pymupdf

    with pymupdf.open(ruta) as doc:
        texto = "\n".join(pagina.get_text() for pagina in doc)
    if not texto.strip():
        raise PdfEscaneadoError(
            f"El PDF no tiene capa de texto (parece escaneado): {ruta}. "
            "El OCR no está soportado en esta versión."
        )
    return texto
